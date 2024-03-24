import os
from tkinter import Tk, filedialog
from docx import Document

# Function definitions
def select_file(title="Select Word File", filetypes=[("Word Files", "*.docx")]):
    root = Tk()
    root.withdraw()  # Hide the main window
    file_path = filedialog.askopenfilename(title=title, filetypes=filetypes)
    return file_path

def read_word_file(file_path):
    try:
        with open(file_path, 'rb') as f:
            document = Document(f)
            text = ""
            for paragraph in document.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
    except Exception as e:
        print("Error reading the Word file:", e)
        return None

def rearrange_text(content):
    # Split the content into lines
    lines = content.strip().split('\n')

    # Initialize an empty list to store word-meaning pairs
    word_meaning_pairs = []

    # Iterate over lines and extract word-meaning pairs
    for line in lines:
        word, meaning = line.split(': ', 1)
        word_meaning_pairs.append((word, meaning))

    # Sort the list alphabetically based on the first item in each tuple
    sorted_pairs = sorted(word_meaning_pairs, key=lambda x: x[0])

    return sorted_pairs

def save_to_doc(table, folder_path):
    doc = Document()
    doc.add_paragraph(table)
    doc_path = os.path.join(folder_path, 'dictionary.docx')
    doc.save(doc_path)
    print(f"Table saved successfully at: {doc_path}")

def select_folder():
    root = Tk()
    root.withdraw()  # Hide the main window
    folder_path = filedialog.askdirectory(title="Select Destination Folder")
    return folder_path

def create_table(sorted_pairs):
    # Find the maximum length of the first column
    max_length = max(len(pair[0]) for pair in sorted_pairs)

    table = 'New Words\t|\tMeaning\n'
    table += '-' * (max_length + 35) + '\n'  # Adjusting the width of the line separator

    for pair in sorted_pairs:
        table += f'{pair[0]:<{max_length}}\t|\t{pair[1]}\n'

    return table


# Run the Program:
file_path = select_file()
if file_path:
    content = read_word_file(file_path)

    sorted_pairs = rearrange_text(content)

    folder_path = select_folder()
    if folder_path:
        table = create_table(sorted_pairs)
        save_to_doc(table, folder_path)
